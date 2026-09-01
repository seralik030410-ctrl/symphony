"""Opt-in real Ollama + Docker acceptance, isolated data; never touches user chats."""
import asyncio
import base64
from io import BytesIO
import json
import os
from pathlib import Path
import time
import uuid

import httpx
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
RUN = ROOT / "data" / "acceptance" / ("stage6-" + uuid.uuid4().hex[:10])
os.environ["SYMPHONY_DATABASE_PATH"] = str(RUN / "symphony.db")
os.environ["SYMPHONY_WORKSPACE_ROOT"] = str(RUN / "workspaces")
os.environ["SYMPHONY_SKILLS_ROOT"] = str(RUN / "skills")
os.environ["SYMPHONY_SEED_BUNDLED_SKILLS"] = "0"

from backend.config import Settings
from backend.main import create_app
from backend.tools.contracts import ToolContext
from backend.tools.context import OcrImageTool, OcrInput


async def main():
    model = os.getenv("SYMPHONY_LIVE_MODEL", "qwen3.5:9b")
    settings = Settings.from_env()
    settings.provider_timeout_seconds = 180
    app = create_app(settings)
    runtime = app.state.runtime
    report = {"model": model, "checks": {}, "ok": False, "directory": str(RUN)}
    async def create(client, title):
        response = await client.post("/api/sessions", json={"title": title, "provider": "ollama", "model": model})
        response.raise_for_status()
        return response.json()
    async def turn(client, session_id, content, **kwargs):
        response = await client.post(f"/api/sessions/{session_id}/turns", json={"content": content, **kwargs})
        response.raise_for_status()
        created = response.json()
        deadline = time.monotonic() + 240
        while time.monotonic() < deadline:
            current = runtime.repository.get_turn(created["turn"]["id"])
            if current["status"] in {"completed", "failed", "cancelled", "interrupted"}:
                events = runtime.repository.list_events(current["id"])
                answer = runtime.repository.get_message(current["assistant_message_id"])["content"]
                assert current["status"] == "completed", current
                return answer, events
            await asyncio.sleep(0.5)
        await runtime.turn_service.cancel(created["turn"]["id"])
        raise AssertionError("Live turn timed out")
    try:
        async with app.router.lifespan_context(app), httpx.AsyncClient(transport=httpx.ASGITransport(app=app), base_url="http://test", timeout=240) as client:
            print("Real Docker extraction and OCR", flush=True)
            session = await create(client, "Stage 6 · Files and OCR")
            from docx import Document
            doc = Document()
            doc.add_paragraph("Calibration reference: Zorabel code is 73421.")
            for _ in range(120):
                doc.add_paragraph("Supporting background material for the test document. " * 15)
            buffer = BytesIO(); doc.save(buffer)
            uploaded = await client.post(f"/api/sessions/{session['id']}/inputs", json={"filename": "reference.docx", "content_base64": base64.b64encode(buffer.getvalue()).decode()})
            uploaded.raise_for_status(); attachment = uploaded.json()
            assert attachment["indexed"]["chunk_count"] > 10
            report["checks"]["isolated_office_index"] = attachment["indexed"]["chunk_count"]
            image = Image.new("RGB", (900, 260), "white")
            draw = ImageDraw.Draw(image)
            font_path = Path("C:/Windows/Fonts/arial.ttf")
            font = ImageFont.truetype(str(font_path), 64) if font_path.exists() else ImageFont.load_default(size=64)
            draw.rectangle((20, 20, 140, 220), fill="blue")
            draw.text((185, 95), "TOTAL 42 USD", fill="black", font=font)
            buffer = BytesIO(); image.save(buffer, "PNG")
            sample = RUN / "ocr-sample.png"; image.save(sample)
            uploaded_image = await client.post(f"/api/sessions/{session['id']}/inputs", json={"filename": "receipt.png", "content_base64": base64.b64encode(buffer.getvalue()).decode()})
            uploaded_image.raise_for_status(); picture = uploaded_image.json()
            turn_record = runtime.repository.create_turn(session["id"], "OCR fixture")["turn"]
            ocr = await OcrImageTool(runtime.file_index, runtime.sandbox).execute(ToolContext(session["id"], turn_record["id"]), OcrInput(path=picture["path"], language="eng"))
            runtime.repository.set_turn_status(turn_record["id"], "completed", finished=True)
            runtime.repository.set_message_status(turn_record["assistant_message_id"], "complete")
            assert "42" in ocr.output["text"], ocr.output
            report["checks"]["local_ocr"] = ocr.output["text"]
            print("Real Ollama retrieval", flush=True)
            answer, events = await turn(client, session["id"], "What is the Zorabel code? Answer with the code from the attached reference.", attachment_ids=[attachment["id"]])
            assert "73421" in answer, answer
            retrieved = next(event for event in events if event["type"] == "context.retrieved")
            assert retrieved["payload"]["characters"] <= 6000
            report["checks"]["retrieval"] = {"answer": answer, "characters": retrieved["payload"]["characters"]}
            print("Real Ollama vision", flush=True)
            caps = await runtime.gateway.resolve_capabilities("ollama", model)
            assert caps.vision, "Choose an installed vision model using SYMPHONY_LIVE_MODEL"
            answer, events = await turn(client, session["id"], "What colour is the rectangle in this image and what number is printed? Reply briefly.", attachment_ids=[picture["id"]])
            assert "42" in answer and ("blue" in answer.lower() or "син" in answer.lower()), answer
            report["checks"]["vision"] = answer
            print("Real Ollama semantic memory", flush=True)
            memory_session = await create(client, "Stage 6 · Semantic memory")
            for index in range(12):
                prompt = ("Background discussion. " * 50 + "Confirmed: project Lumen uses SQLite; network must stay disabled. TODO: verify the release.") if index == 0 else f"Discussion {index}: continue the Lumen project without changing earlier decisions."
                item = runtime.repository.create_turn(memory_session["id"], prompt)
                runtime.repository.append_assistant_delta(item["assistant_message"]["id"], "Acknowledged. No files changed and release verification is still pending.")
                runtime.repository.set_message_status(item["assistant_message"]["id"], "complete")
                runtime.repository.set_turn_status(item["turn"]["id"], "completed", finished=True)
            response = await client.post(f"/api/sessions/{memory_session['id']}/memory/snapshot")
            response.raise_for_status(); memory = response.json()
            serialized = json.dumps(memory, ensure_ascii=False)
            assert "Lumen" in serialized and "SQLite" in serialized and memory["open_tasks"], memory
            assert len(memory["source_message_ids"]) == 14
            report["checks"]["semantic_memory"] = memory
            clean = await create(client, "Stage 6 · Isolation")
            assert runtime.memory.get(clean["id"])["version"] == 0 and runtime.file_index.list_files(clean["id"]) == []
            answer, events = await turn(client, clean["id"], "Почему небо голубое? Ответь одним предложением.")
            assert not any(event["type"].startswith("tool.") or event["type"] == "context.retrieved" for event in events)
            report["checks"]["ordinary_chat"] = answer
            report["ok"] = True
    finally:
        RUN.mkdir(parents=True, exist_ok=True)
        (RUN / "report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        print(json.dumps({"ok": report["ok"], "report": str(RUN / "report.json")}), flush=True)


if __name__ == "__main__":
    asyncio.run(main())
